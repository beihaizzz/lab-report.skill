#!/usr/bin/env python3
"""生成通用实验报告模板（无敏感信息，仅含占位符）。"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _set_cell(cell, text, role="值", font_size=14):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.size = Pt(font_size)

    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    if role == "标题":
        rFonts.set(qn('w:eastAsia'), "\u9ed1\u4f53")  # 黑体
    else:
        run.font.name = "\u5b8b\u4f53"  # 宋体
    rPr.append(rFonts)


def _add_section_heading(doc, text, font_size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), "\u9ed1\u4f53")
    rPr.append(rFonts)
    return p


def main():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # ===== 封面 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{学院名称}}")
    run.font.size = Pt(36)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("{{报告标题}}")
    run2.font.size = Pt(36)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("（ {{学年学期}} ）")
    run3.font.size = Pt(16)

    doc.add_paragraph()

    # ===== Table 0：学生信息 =====
    t0 = doc.add_table(rows=6, cols=4)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_def = [
        [("课程名称", "{{课程名称}}", 2), (None, None, 0), (None, None, 0)],
        [("课程代码", "{{课程代码}}", 2), (None, None, 0), (None, None, 0)],
        [("任课教师", "{{任课教师}}", 2), (None, None, 0), (None, None, 0)],
        [("学生姓名", "{{学生姓名}}", 1), ("专业年级", "{{专业年级}}", 1)],
        [("学号", "{{学号}}", 1), ("实验成绩", "{{实验成绩}}", 1)],
        [("实验学时", "{{实验学时}}", 1), ("实验日期", "{{实验日期}}", 1)],
    ]

    for ri, row_data in enumerate(rows_def):
        col_pairs = [(i, d) for i, d in enumerate(row_data) if d[0] is not None]
        for pair_idx, (orig_col, (label, value, colspan)) in enumerate(col_pairs):
            # 标题列
            label_cell = t0.cell(ri, pair_idx * 2)
            _set_cell(label_cell, label, "标题")
            # 值列 - 如果 colspan>1 需要合并
            val_cell = t0.cell(ri, pair_idx * 2 + 1)
            if colspan > 1:
                # 合并后续单元格
                for extra in range(1, colspan):
                    target_cell = t0.cell(ri, pair_idx * 2 + 1 + extra)
                    val_cell.merge(target_cell)
            _set_cell(val_cell, value, "值")

    doc.add_paragraph()

    # ===== Table 1：实验信息 =====
    t1 = doc.add_table(rows=3, cols=4)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER

    t1_rows = [
        [("实验名称", "{{实验名称}}"), ("实验类型", "{{实验类型}}")],
        [("实验地点", "{{实验地点}}"), ("实验环境", "{{实验环境}}")],
        [("实验设备", "{{实验设备}}"), ("提交文档", "{{提交文档说明}}")],
    ]

    for ri, row_data in enumerate(t1_rows):
        for pair_idx, (label, value) in enumerate(row_data):
            label_cell = t1.cell(ri, pair_idx * 2)
            _set_cell(label_cell, label, "标题", font_size=12)
            val_cell = t1.cell(ri, pair_idx * 2 + 1)
            _set_cell(val_cell, value, "值", font_size=12)

    doc.add_paragraph()

    # ===== 内容区域 ====="
    sections = [
        ("h1", "实验目的"),
        ("placeholder", "{{实验目的}}"),
        ("h1", "实验要求"),
        ("placeholder", "{{实验要求}}"),
        ("h1", "实验原理"),
        ("placeholder", "{{实验原理}}"),
        ("h1", "实验内容"),
        ("placeholder", "{{实验内容}}"),
        ("h1", "实验总结"),
        ("placeholder", "{{实验总结}}"),
        ("h1", "附：源代码"),
        ("placeholder", "{{源代码}}"),
    ]

    for role, text in sections:
        if role == "h1":
            _add_section_heading(doc, text)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            run = p.add_run(text)
            run.font.name = "\u5b8b\u4f53"
            run.font.size = Pt(12)

    output_path = "E:/lab-report/lab-report/assets/report_template.docx"
    doc.save(output_path)
    print(f"模板已保存: {output_path}")


if __name__ == "__main__":
    main()
