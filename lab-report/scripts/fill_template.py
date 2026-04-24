#!/usr/bin/env python3
"""Template filling with docxtpl, CJK support, de-AI styles, and format inheritance."""

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docxtpl import DocxTemplate
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

BANNED_WORDS = ['首先', '其次', '最后', '总而言之', '值得注意的是', '综上所述', '不可否认']

# ── Style map: map placeholder names to formatting roles ────────────────────
HEADING_PLACEHOLDERS = {
    '实验目的', '实验原理', '实验器材', '实验步骤',
    '实验数据', '实验结果', '实验结论', '实验要求',
}
CODE_PLACEHOLDERS = {'实验程序', '程序代码', '代码', '实验代码'}
INFO_PLACEHOLDERS = {'姓名', '学号', '学院', '专业', '班级', '课程名',
                     '实验名称', '实验日期', '实验地点'}

STYLE_MAP = {
    'title':     {'font_name': '黑体', 'font_size': Pt(12), 'bold': True, 'align': WD_PARAGRAPH_ALIGNMENT.CENTER},
    'heading1':  {'font_name': '黑体', 'font_size': Pt(12), 'bold': True, 'align': WD_PARAGRAPH_ALIGNMENT.LEFT},
    'heading2':  {'font_name': '黑体', 'font_size': Pt(11), 'bold': True, 'align': WD_PARAGRAPH_ALIGNMENT.LEFT},
    'body':      {'font_name': '宋体', 'font_size': Pt(10.5), 'bold': False, 'align': None},
    'code':      {'font_name': 'Courier New', 'font_size': Pt(9), 'bold': False, 'align': None},
    'info_cell': {'font_name': '宋体', 'font_size': Pt(10.5), 'bold': False, 'align': WD_PARAGRAPH_ALIGNMENT.CENTER},
}


def _find_libreoffice() -> str | None:
    for name in ['soffice', 'libreoffice', 'soffice.exe', 'libreoffice.exe']:
        found = shutil.which(name)
        if found:
            return found
    for base in [r'C:\Program Files\LibreOffice\program',
                 r'C:\Program Files (x86)\LibreOffice\program']:
        for exe in ['soffice.exe', 'swriter.exe']:
            cand = Path(base) / exe
            if cand.exists():
                return str(cand)
    return None


def _convert_to_docx(filepath: Path) -> Path | None:
    lo = _find_libreoffice()
    if not lo:
        return None
    try:
        result = subprocess.run(
            [lo, '--headless', '--convert-to', 'docx', '--outdir',
             str(filepath.parent), str(filepath)],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            docx = filepath.with_suffix('.docx')
            if docx.exists():
                return docx
    except Exception:
        pass
    return None


def set_cjk_font(run, font_name='宋体'):
    """Set CJK font for a run to prevent tofu."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def apply_style_to_run(run, style: dict):
    """Apply a STYLE_MAP dict to a run."""
    run.font.name = style.get('font_name', '宋体')
    set_cjk_font(run, style.get('font_name', '宋体'))
    if style.get('font_size'):
        run.font.size = style['font_size']
    if style.get('bold'):
        run.font.bold = True


def apply_style_to_paragraph(para, style: dict):
    """Apply a STYLE_MAP dict to a paragraph (alignment + runs)."""
    align = style.get('align')
    if align is not None:
        para.alignment = align
    for run in para.runs:
        apply_style_to_run(run, style)


def _detect_placeholder_role(text: str) -> str:
    """Detect the role of a placeholder from its content."""
    content = text.strip('{} \t\n\r')
    if content in HEADING_PLACEHOLDERS:
        return 'heading1'
    if content in CODE_PLACEHOLDERS:
        return 'code'
    if content in INFO_PLACEHOLDERS:
        return 'info_cell'
    return 'body'


def _snapshot_cell_styles(doc: Document) -> dict:
    """Snapshot original cell formatting (font/size/bold/align) for each cell index."""
    snap = {}
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    for r in p.runs:
                        key = (table_idx, row_idx, col_idx)
                        snap.setdefault(key, {
                            'font_name': r.font.name or getattr(r.font, 'name', None) or '宋体',
                            'font_size': r.font.size,
                            'bold': r.font.bold,
                            'align': p.alignment,
                        })
    return snap


def _apply_cell_styles_from_snapshot(cell, style_ref: dict):
    """Apply formatting from a snapshot dict to a cell's content."""
    for p in cell.paragraphs:
        if style_ref.get('align') is not None:
            p.alignment = style_ref['align']
        for r in p.runs:
            if style_ref.get('font_name'):
                r.font.name = style_ref['font_name']
                set_cjk_font(r, style_ref['font_name'])
            if style_ref.get('font_size'):
                r.font.size = style_ref['font_size']
            if style_ref.get('bold'):
                r.font.bold = True


def _post_process_formatting(doc: Document, style_snap: dict):
    """Post-process: apply styles based on content role and snapshot data."""
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue

                # 1) Try to apply snapshot style (inheritance from template)
                snap_key = (table_idx, row_idx, col_idx)
                if snap_key in style_snap:
                    _apply_cell_styles_from_snapshot(cell, style_snap[snap_key])
                else:
                    # 2) Fallback: detect role from content
                    role = _detect_placeholder_role(cell_text)
                    style = STYLE_MAP.get(role, STYLE_MAP['body'])
                    apply_style_to_paragraph(cell.paragraphs[0], style)

                # 3) Normalize: remove extra empty paragraphs (fix 2.3)
                _remove_extra_empty_paragraphs(cell)

    # Apply style mapping to body paragraphs too
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        role = _detect_placeholder_role(text)
        style = STYLE_MAP.get(role, STYLE_MAP['body'])
        apply_style_to_paragraph(para, style)


def _remove_extra_empty_paragraphs(container):
    """Remove excessive empty paragraphs from a cell/body (fix 2.3)."""
    paragraphs = container.paragraphs
    # If more than 1 para and some are empty, collapse consecutive empties
    if len(paragraphs) <= 1:
        return
    i = 0
    while i < len(paragraphs):
        if i > 0 and not paragraphs[i].text.strip() and not paragraphs[i - 1].text.strip():
            # Found consecutive empty paras — remove this one
            p_element = paragraphs[i]._element
            p_element.getparent().remove(p_element)
            # Don't increment i; the list shifted
        else:
            i += 1


def _insert_image_placeholders(doc: Document, image_instructions: list):
    """Insert image placeholder markers where requested.
    
    image_instructions: list of dicts with keys:
      - placeholder: text marker in the template (e.g. '[insert_image_接线图]')
      - label: human-readable label (e.g. '请在此处粘贴：硬件接线照片')
    """
    for para in doc.paragraphs:
        for instr in image_instructions:
            if instr['placeholder'] in para.text:
                para.clear()
                run = para.add_run(f"\n[{instr['label']}]\n")
                run.font.name = '宋体'
                run.font.size = Pt(10)
                run.font.color.rgb = None  # keep default
                run.italic = True
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def fill_template(template_path: Path, data_path: Path, output_path: Path,
                  style: str = "normal", image_placeholders: list = None):
    """Fill template with data."""
    if not HAS_DOCX:
        return {"error": "Missing dependencies: python-docx, docxtpl"}

    result = {
        "success": False,
        "template": str(template_path),
        "output": str(output_path),
        "placeholders_filled": [],
        "placeholders_missing": [],
        "style_applied": style,
        "formatting_warnings": [],
    }

    try:
        # Load data
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Handle .doc → .docx conversion
        suffix = template_path.suffix.lower()
        if suffix == '.doc':
            converted = _convert_to_docx(template_path)
            if converted is None:
                result["error"] = (
                    "Cannot process .doc template. Please:\n"
                    "1. Install LibreOffice (https://www.libreoffice.org/)\n"
                    "2. Or manually save as .docx in Word / WPS"
                )
                return result
            template_path = converted

        # ═══ Phase 1: Snapshot template cell styles ═══
        temp_doc = Document(template_path)
        style_snap = _snapshot_cell_styles(temp_doc)

        # ═══ Phase 2: Render via docxtpl ═══
        shutil.copy(template_path, output_path)
        doc = DocxTemplate(output_path)
        doc.render(data)
        doc.save(output_path)

        # ═══ Phase 3: Post-process formatting ═══
        final_doc = Document(output_path)

        # 3a — CJK font fix
        for para in final_doc.paragraphs:
            for run in para.runs:
                if any('\u4e00' <= c <= '\u9fff' for c in run.text):
                    set_cjk_font(run)

        # 3b — Style inheritance + heading/body mapping + empty para removal
        _post_process_formatting(final_doc, style_snap)

        # 3c — Image placeholders (2.4)
        if image_placeholders:
            _insert_image_placeholders(final_doc, image_placeholders)

        final_doc.save(output_path)

        # ═══ Phase 4: Verify ═══
        result["success"] = True
        verify_doc = Document(output_path)
        full_text = " ".join([p.text for p in verify_doc.paragraphs])
        remaining = re.findall(r'\{\{([^}]+)\}\}', full_text)
        result["placeholders_missing"] = remaining

    except Exception as e:
        result["error"] = str(e)

    return result


def verify_original_unchanged(template_path: Path, original_hash: str) -> bool:
    with open(template_path, 'rb') as f:
        current_hash = hashlib.sha256(f.read()).hexdigest()
    return current_hash == original_hash


def main():
    parser = argparse.ArgumentParser(description='Fill DOCX template')
    parser.add_argument('--template', '-t', required=True, help='Template DOCX file')
    parser.add_argument('--data', '-d', required=True, help='JSON data file')
    parser.add_argument('--output', '-o', required=True, help='Output DOCX file')
    parser.add_argument('--style', '-s', choices=['perfect', 'normal'], default='normal',
                        help='normal=标准报告(90+分, 日常首选); perfect=极尽详尽(特殊场景)')
    parser.add_argument('--image-placeholders', help='JSON file with image placeholder definitions')
    args = parser.parse_args()

    template_path = Path(args.template)
    data_path = Path(args.data)
    output_path = Path(args.output)

    if not template_path.exists():
        print(f"Error: Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)
    if not data_path.exists():
        print(f"Error: Data file not found: {data_path}", file=sys.stderr)
        sys.exit(1)

    with open(template_path, 'rb') as f:
        original_hash = hashlib.sha256(f.read()).hexdigest()

    # Load image placeholder instructions
    image_placeholders = None
    if args.image_placeholders:
        ip_path = Path(args.image_placeholders)
        if ip_path.exists():
            with open(ip_path, 'r', encoding='utf-8') as f:
                image_placeholders = json.load(f)

    result = fill_template(template_path, data_path, output_path, args.style, image_placeholders)

    if not verify_original_unchanged(template_path, original_hash):
        print("Error: Original template was modified!", file=sys.stderr)
        sys.exit(1)

    print(json.dumps(result, indent=2, ensure_ascii=False))
    sys.exit(0 if result.get("success") else 1)


if __name__ == '__main__':
    main()
